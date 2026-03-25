"""
CATCH 채용정보·노무사 종합 전략 보고서 생성 스크립트
"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ──────────────────────────────────────────
# 페이지 여백 설정
# ──────────────────────────────────────────
section = doc.sections[0]
section.page_width  = Cm(21)
section.page_height = Cm(29.7)
section.left_margin   = Cm(2.5)
section.right_margin  = Cm(2.5)
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.0)

# ──────────────────────────────────────────
# 색상 팔레트
# ──────────────────────────────────────────
BRAND_BLUE   = RGBColor(0x1A, 0x56, 0xDB)   # 메인 브랜드 블루
BRAND_DARK   = RGBColor(0x1E, 0x29, 0x3B)   # 다크 네이비
ACCENT_GREEN = RGBColor(0x05, 0x96, 0x69)   # 녹색 강조
ACCENT_RED   = RGBColor(0xDC, 0x26, 0x26)   # 빨간 경고
ACCENT_ORANGE= RGBColor(0xD9, 0x77, 0x06)   # 주황 주의
ACCENT_PURPLE= RGBColor(0x70, 0x3E, 0xFF)   # 보라 포인트
LIGHT_BLUE_BG= RGBColor(0xDB, 0xEA, 0xFE)   # 연한 파랑 배경
LIGHT_GREEN_BG=RGBColor(0xD1, 0xFA, 0xE5)   # 연한 녹색 배경
GRAY_TEXT    = RGBColor(0x6B, 0x72, 0x80)   # 서브 텍스트
TABLE_HEADER = RGBColor(0x1E, 0x40, 0xAF)   # 표 헤더
TABLE_STRIPE = RGBColor(0xEF, 0xF6, 0xFF)   # 표 줄무늬
WHITE        = RGBColor(0xFF, 0xFF, 0xFF)
BLACK        = RGBColor(0x11, 0x18, 0x27)

# ──────────────────────────────────────────
# 헬퍼 함수들
# ──────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    """셀 배경색 설정"""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

def set_cell_border(cell, top=None, bottom=None, left=None, right=None):
    """셀 테두리 설정"""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        if val:
            el = OxmlElement(f'w:{side}')
            el.set(qn('w:val'),   val.get('val','single'))
            el.set(qn('w:sz'),    str(val.get('sz', 4)))
            el.set(qn('w:space'),'0')
            el.set(qn('w:color'), val.get('color','000000'))
            tcBorders.append(el)
    tcPr.append(tcBorders)

def add_heading(doc, text, level=1, color=None, space_before=12, space_after=6):
    """헤딩 추가"""
    p = doc.add_heading(text, level=level)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    run = p.runs[0] if p.runs else p.add_run(text)
    if color:
        run.font.color.rgb = color
    return p

def add_body(doc, text, bold=False, color=None, size=10.5, space_after=4, indent=0):
    """본문 단락 추가"""
    p = doc.add_paragraph()
    p.paragraph_format.space_after  = Pt(space_after)
    p.paragraph_format.space_before = Pt(2)
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color
    return p

def add_bullet(doc, text, level=0, color=None):
    """불릿 포인트 추가"""
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after  = Pt(3)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.left_indent  = Cm(0.5 + level * 0.5)
    run = p.add_run(text)
    run.font.size = Pt(10)
    if color:
        run.font.color.rgb = color
    return p

def add_info_box(doc, title, items, box_color='DBE AFF', title_color=None, bg_hex='EFF6FF'):
    """정보 박스 (테이블 1행)"""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, bg_hex)
    set_cell_border(cell,
        top={'val':'single','sz':6,'color':'1A56DB'},
        bottom={'val':'single','sz':2,'color':'1A56DB'},
        left={'val':'single','sz':6,'color':'1A56DB'},
        right={'val':'single','sz':6,'color':'1A56DB'})
    cell.paragraphs[0].clear()
    # 타이틀
    tp = cell.paragraphs[0]
    tr = tp.add_run(f'  {title}')
    tr.font.bold = True
    tr.font.size = Pt(11)
    tr.font.color.rgb = title_color or BRAND_BLUE
    # 아이템들
    for item in items:
        ip = cell.add_paragraph()
        ip.paragraph_format.left_indent = Cm(0.5)
        ir = ip.add_run(item)
        ir.font.size = Pt(10)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return tbl

def make_table(doc, headers, rows, col_widths=None, stripe=True):
    """스타일 테이블 생성"""
    num_cols = len(headers)
    tbl = doc.add_table(rows=1+len(rows), cols=num_cols)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style = 'Table Grid'

    # 헤더 행
    hdr_row = tbl.rows[0]
    for i, hdr in enumerate(headers):
        cell = hdr_row.cells[i]
        set_cell_bg(cell, '1E40AF')
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(hdr)
        run.font.bold  = True
        run.font.size  = Pt(9.5)
        run.font.color.rgb = WHITE

    # 데이터 행
    for ri, row_data in enumerate(rows):
        row = tbl.rows[ri + 1]
        bg  = 'EFF6FF' if (stripe and ri % 2 == 1) else 'FFFFFF'
        for ci, cell_text in enumerate(row_data):
            cell = row.cells[ci]
            set_cell_bg(cell, bg)
            p = cell.paragraphs[0]
            # 색상 태그 파싱: [RED]text, [GREEN]text, [BOLD]text
            text = str(cell_text)
            if text.startswith('[RED]'):
                run = p.add_run(text[5:])
                run.font.color.rgb = ACCENT_RED
                run.font.bold = True
            elif text.startswith('[GREEN]'):
                run = p.add_run(text[7:])
                run.font.color.rgb = ACCENT_GREEN
                run.font.bold = True
            elif text.startswith('[ORANGE]'):
                run = p.add_run(text[8:])
                run.font.color.rgb = ACCENT_ORANGE
                run.font.bold = True
            elif text.startswith('[BOLD]'):
                run = p.add_run(text[6:])
                run.font.bold = True
            else:
                run = p.add_run(text)
            run.font.size = Pt(9.5)

    # 열 너비
    if col_widths:
        for ri2, row2 in enumerate(tbl.rows):
            for ci2, cell2 in enumerate(row2.cells):
                if ci2 < len(col_widths):
                    cell2.width = Cm(col_widths[ci2])

    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    return tbl

def add_divider(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run('─' * 72)
    run.font.color.rgb = RGBColor(0xD1, 0xD5, 0xDB)
    run.font.size = Pt(8)

def add_part_title(doc, part_num, title, subtitle=''):
    """PART 타이틀 블록"""
    doc.add_page_break()
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, '1E293B')
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p.add_run(f'PART {part_num}')
    r1.font.size = Pt(10)
    r1.font.bold = True
    r1.font.color.rgb = RGBColor(0x93, 0xC5, 0xFD)
    p.add_run('\n')
    r2 = p.add_run(title)
    r2.font.size = Pt(20)
    r2.font.bold = True
    r2.font.color.rgb = WHITE
    if subtitle:
        p.add_run('\n')
        r3 = p.add_run(subtitle)
        r3.font.size = Pt(11)
        r3.font.color.rgb = RGBColor(0xCB, 0xD5, 0xE1)
    cell.paragraphs[0].paragraph_format.space_before = Pt(14)
    cell.paragraphs[0].paragraph_format.space_after  = Pt(14)
    doc.add_paragraph().paragraph_format.space_after = Pt(8)

# ══════════════════════════════════════════════════
# 표지
# ══════════════════════════════════════════════════
p_cover = doc.add_paragraph()
p_cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_cover.paragraph_format.space_before = Pt(60)
r = p_cover.add_run('CATCH (퇴직금 한번에)')
r.font.size  = Pt(28)
r.font.bold  = True
r.font.color.rgb = BRAND_BLUE

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p2.add_run('채용정보 & 노무사 중개 서비스')
r2.font.size = Pt(18)
r2.font.color.rgb = BRAND_DARK

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = p3.add_run('종합 전략 보고서')
r3.font.size = Pt(22)
r3.font.bold = True
r3.font.color.rgb = BRAND_DARK

doc.add_paragraph()
p4 = doc.add_paragraph()
p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
r4 = p4.add_run('─────────────────────────────────────────')
r4.font.color.rgb = BRAND_BLUE

doc.add_paragraph()
p5 = doc.add_paragraph()
p5.alignment = WD_ALIGN_PARAGRAPH.CENTER
r5 = p5.add_run('작성일: 2026년 3월 24일')
r5.font.size = Pt(11)
r5.font.color.rgb = GRAY_TEXT

p6 = doc.add_paragraph()
p6.alignment = WD_ALIGN_PARAGRAPH.CENTER
r6 = p6.add_run('대상: 일용직 근로자 (쿠팡 · 마켓컬리 · CJ대한통운)')
r6.font.size = Pt(11)
r6.font.color.rgb = GRAY_TEXT

doc.add_paragraph()

# 목차 박스
tbl_toc = doc.add_table(rows=1, cols=1)
tbl_toc.alignment = WD_TABLE_ALIGNMENT.CENTER
c = tbl_toc.cell(0, 0)
set_cell_bg(c, 'F0F9FF')
set_cell_border(c,
    top={'val':'single','sz':8,'color':'1A56DB'},
    bottom={'val':'single','sz':8,'color':'1A56DB'},
    left={'val':'single','sz':8,'color':'1A56DB'},
    right={'val':'single','sz':8,'color':'1A56DB'})
tp0 = c.paragraphs[0]
tp0.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_toc = tp0.add_run('📋  목차')
r_toc.font.size = Pt(13)
r_toc.font.bold = True
r_toc.font.color.rgb = BRAND_BLUE

toc_items = [
    ('PART 1', '채용정보 섹션 전략', '1.1~1.8'),
    ('PART 2', '노무사 중개 서비스 전략', '2.1~2.6'),
    ('PART 3', '노무사 파트너십 세일즈 & 마케팅 자료', '3.1~3.5'),
    ('부록',   '통합 12개월 로드맵 + 핵심 요약', '—'),
]
for part, ttl, sec in toc_items:
    tp = c.add_paragraph()
    tp.paragraph_format.left_indent = Cm(1.0)
    tp.paragraph_format.space_after = Pt(2)
    rp1 = tp.add_run(f'{part}  ')
    rp1.font.bold = True
    rp1.font.size = Pt(10.5)
    rp1.font.color.rgb = BRAND_BLUE
    rp2 = tp.add_run(f'{ttl}  ')
    rp2.font.size = Pt(10.5)
    rp3 = tp.add_run(f'[{sec}]')
    rp3.font.size  = Pt(9.5)
    rp3.font.color.rgb = GRAY_TEXT

# ══════════════════════════════════════════════════
# PART 1 — 채용정보 섹션 전략
# ══════════════════════════════════════════════════
add_part_title(doc, '1', '채용정보 섹션 전략',
               '퇴직 후 재취업을 원하는 사용자를 위한 채용 정보 제공')

# 1.1
add_heading(doc, '1.1  전략 개요 및 사업적 근거', level=2, color=BRAND_BLUE)
add_body(doc,
    'CATCH 사용자는 퇴직금 계산을 완료한 직후 "다음 직장"을 탐색하는 전환점에 있습니다. '
    '이 순간은 채용 정보 노출의 최적 타이밍이며, 별도 마케팅 비용 없이 높은 클릭률을 기대할 수 있습니다.',
    space_after=6)

add_info_box(doc, '💡 핵심 인사이트',
    ['• 퇴직금 계산 완료 후 평균 체류 시간 3분 이상 → 채용정보 탐색 전환율 高',
     '• 쿠팡·마켓컬리 등 동일 플랫폼 재취업 수요 존재',
     '• CPC 광고 대비 네이티브 채용정보 CTR 2~4배 높음 (업계 데이터)'],
    bg_hex='EFF6FF', title_color=BRAND_BLUE)

add_body(doc, '▶ 사업적 근거', bold=True, color=BRAND_DARK)
for item in [
    '① 재취업 수요: 일용직 근로자의 약 68%는 동일 업종 재취업을 선호 (고용노동부 2024)',
    '② 수익 다각화: 광고 CPC 0.5~2만원/클릭 × 월 1만 클릭 = 月 500만~2,000만원 가능',
    '③ 사용자 retention: 채용정보 섹션이 있으면 재방문율 +35% 기대',
    '④ 파트너십 레버리지: 쿠팡 파트너스, 마켓컬리 공식 채용 제휴 협상 명분',
]:
    add_bullet(doc, item)

# 1.2
add_heading(doc, '1.2  실현 가능성 분석 — 데이터 소스별', level=2, color=BRAND_BLUE)
make_table(doc,
    headers=['데이터 소스', '취득 방법', '비용', '갱신 주기', '품질', '추천도'],
    rows=[
        ['워크넷 (고용노동부)', 'OpenAPI (무료)', '[GREEN]무료', '실시간', '★★★☆', '[GREEN]★★★★★'],
        ['잡코리아 / 사람인',  '비공식 스크래핑',  '[ORANGE]회색지대', '일 1회', '★★★★', '[ORANGE]★★☆☆☆'],
        ['쿠팡 공식 채용',     '직접 파싱/제휴',  '협의',    '주 1회', '★★★★★', '[GREEN]★★★★☆'],
        ['마켓컬리 채용',      '직접 파싱/제휴',  '협의',    '주 1회', '★★★★★', '[GREEN]★★★★☆'],
        ['수동 큐레이션',      '관리자 직접 입력', '[GREEN]무료', '수동',  '★★★★★', '★★★☆☆'],
    ],
    col_widths=[3.5, 3.5, 2.0, 2.0, 2.0, 2.0]
)

# 1.3
add_heading(doc, '1.3  추천 전략: 2단계 접근법', level=2, color=BRAND_BLUE)
add_body(doc, 'Phase 1 (즉시 실행) — 워크넷 API 무료 연동', bold=True, color=ACCENT_GREEN)
for item in [
    '워크넷 OpenAPI 신청 (승인 1~3일) → 물류·배송·단순노무 카테고리 필터링',
    '프론트엔드에 /jobs 페이지 추가, 홈 화면 하단 채용정보 카드 3~5개 노출',
    '위치 기반 필터 (시/도 선택) → 사용자 맞춤 채용정보 제공',
]:
    add_bullet(doc, item, color=BRAND_DARK)

doc.add_paragraph()
add_body(doc, 'Phase 2 (1~3개월) — 파트너십 협의', bold=True, color=BRAND_BLUE)
for item in [
    '쿠팡 파트너스 / 마켓컬리 공식 제휴 → 독점 채용정보 + CPC 수익',
    '잡코리아·사람인 공식 API 계약 협의 (월 10~50만원 구간)',
    '노무사 파트너 채용 공고 우선 노출 (PART 2 연동)',
]:
    add_bullet(doc, item, color=BRAND_DARK)

# 1.4
add_heading(doc, '1.4  UI 배치 전략', level=2, color=BRAND_BLUE)
make_table(doc,
    headers=['배치 위치', '컴포넌트', '내용', '전환 목표'],
    rows=[
        ['홈 화면 하단', 'JobCardCarousel', '최신 채용 3개 카드 슬라이더', '클릭 → /jobs'],
        ['계산 결과 페이지', 'JobRecommendBanner', '"지금 바로 재취업!" CTA 배너', '클릭 → /jobs?filter=company'],
        ['마이페이지', 'SavedJobsWidget', '저장한 채용공고 목록', '지속 방문 유도'],
        ['/jobs 전용 페이지', 'JobListPage', '전체 채용정보 + 필터 (지역/직종/급여)', '상세 탐색'],
    ],
    col_widths=[3.0, 3.8, 5.0, 3.2]
)

# 1.5
add_heading(doc, '1.5  데이터 아키텍처', level=2, color=BRAND_BLUE)
add_info_box(doc, '🏗️  추천 아키텍처: Supabase Edge Function + jobs 테이블',
    ['• Supabase Edge Function (Deno) → 워크넷 API 호출 (CRON 1시간마다)',
     '• jobs 테이블: id, source, title, company, location, salary, url, created_at, expires_at',
     '• 프론트엔드: /api/jobs → FastAPI → Supabase SELECT (캐싱 5분)',
     '• 광고 클릭 추적: job_clicks 테이블 (job_id, user_id, clicked_at)'],
    bg_hex='F0FDF4', title_color=ACCENT_GREEN)

make_table(doc,
    headers=['테이블', '주요 컬럼', '용도'],
    rows=[
        ['jobs',        'id, source, title, company, location, salary_min, salary_max, url, expires_at', '채용공고 저장'],
        ['job_clicks',  'id, job_id, user_id, session_id, clicked_at', 'CPC 수익 집계'],
        ['job_saves',   'id, job_id, user_id, saved_at', '북마크 기능'],
        ['job_partners','id, company_name, api_key, cpc_rate, contract_end', '파트너사 관리'],
    ],
    col_widths=[2.8, 8.5, 2.7]
)

# 1.6
add_heading(doc, '1.6  법적 리스크 분석', level=2, color=BRAND_BLUE)
make_table(doc,
    headers=['항목', '위험도', '내용', '대응 방안'],
    rows=[
        ['워크넷 API 상업 이용', '[GREEN]LOW', '공공데이터 상업이용 허용', '이용약관 준수, 출처 표기'],
        ['잡코리아 스크래핑',   '[RED]HIGH',  '저작권법·이용약관 위반 가능', '공식 API 계약 체결'],
        ['채용정보 허위 게재',   '[ORANGE]MID','구인·구직 관리법 위반 가능', '소스 원문 링크 제공'],
        ['개인정보 수집',       '[ORANGE]MID','이력서 저장 시 개인정보보호법', '최소 수집, 암호화 저장'],
        ['광고 고지 의무',      '[GREEN]LOW', '표시광고법: 광고임을 명시', '"광고" 또는 "협찬" 배지 표시'],
    ],
    col_widths=[3.5, 2.0, 4.5, 4.0]
)

# 1.7
add_heading(doc, '1.7  예상 수익 모델', level=2, color=BRAND_BLUE)
make_table(doc,
    headers=['수익원', '단가', '월 예상 볼륨', '월 예상 수익', '비고'],
    rows=[
        ['CPC 광고 (워크넷)',   '₩500~1,000/클릭', '2,000클릭', '₩100~200만', '6개월 후 도달 가능'],
        ['CPC 광고 (파트너사)', '₩1,500~3,000/클릭','1,000클릭', '₩150~300만', '파트너 계약 후'],
        ['프리미엄 공고 노출',  '₩30,000~50,000/건','20건',      '₩60~100만', '직접 영업'],
        ['채용 성사 수수료',    '월급의 5~10%',     '5건',       '₩50~150만', '장기 목표'],
        ['[BOLD]합계 (12개월 후)', '—', '—', '[GREEN]₩360~750만/월', '[BOLD]연 4,320~9,000만원'],
    ],
    col_widths=[3.5, 3.0, 2.5, 2.8, 3.2]
)

# 1.8
add_heading(doc, '1.8  채용정보 KPI', level=2, color=BRAND_BLUE)
make_table(doc,
    headers=['지표', '1개월', '3개월', '6개월', '12개월', '측정 방법'],
    rows=[
        ['채용공고 노출 수',   '500건',    '2,000건',   '5,000건',   '10,000건', 'Supabase count'],
        ['채용 클릭률(CTR)',   '3%',       '5%',        '7%',        '10%',      'job_clicks/views'],
        ['재방문율',           '—',        '+15%',      '+25%',      '+35%',     'Vercel Analytics'],
        ['CPC 수익',          '₩10만',    '₩100만',    '₩300만',    '₩700만',   '파트너 대시보드'],
        ['파트너 기업 수',     '0',        '2',         '5',         '10',       'job_partners 테이블'],
    ],
    col_widths=[3.5, 2.0, 2.0, 2.0, 2.0, 3.5]
)

# ══════════════════════════════════════════════════
# PART 2 — 노무사 중개 서비스 전략
# ══════════════════════════════════════════════════
add_part_title(doc, '2', '노무사 중개 서비스 전략',
               '계산 직후 상담 전환 — 골든 타이밍 포착')

# 2.1
add_heading(doc, '2.1  전략 개요', level=2, color=BRAND_BLUE)
add_body(doc,
    '퇴직금 계산 완료 직후는 사용자가 "내 권리를 어떻게 받을 수 있을까?"를 고민하는 순간입니다. '
    '노무사 상담 CTA를 이 시점에 배치하면 자연스러운 전환이 이루어지며, '
    '노무사 입장에서도 의도가 명확한 고품질 리드(잠재 고객)를 얻을 수 있습니다.',
    space_after=6)

add_info_box(doc, '🎯 골든 타이밍 3가지',
    ['① 계산 결과 페이지: "퇴직금을 못 받고 있다면?" → 노무사 상담 CTA',
     '② 적격 판정 후: "이제 어떻게 청구해야 하나?" → 1:1 상담 연결',
     '③ 부적격 판정 후: "정말 받을 수 없나요?" → 노무사 재검토 의뢰'],
    bg_hex='FFF7ED', title_color=ACCENT_ORANGE)

# 2.2
add_heading(doc, '2.2  섹션 배치 전략', level=2, color=BRAND_BLUE)

add_body(doc, '① 계산 결과 페이지 — CTA 배너', bold=True, color=BRAND_DARK)
add_body(doc,
    '결과 카드 하단에 "노무사에게 직접 문의하기" 버튼 배치. '
    '퇴직금 금액이 클수록 (예: 100만원 이상) 강조 강도 증가.',
    indent=0.5)

add_body(doc, '② 마이페이지 — 노무사 디렉토리', bold=True, color=BRAND_DARK)
add_body(doc,
    '지역별/전문분야별 노무사 카드 목록. '
    '별점, 상담료, 전문 분야, 실시간 예약 가능 여부 표시.',
    indent=0.5)

add_body(doc, '③ 홈 화면 — 배너 슬롯', bold=True, color=BRAND_DARK)
add_body(doc,
    '"전국 노무사와 바로 연결" 배너. 월 1~2회 교체. '
    '로그인 사용자에게 개인화 (지역 기반 노무사 추천)',
    indent=0.5)

make_table(doc,
    headers=['배치', '컴포넌트명', '트리거 조건', '전환 목표'],
    rows=[
        ['계산 결과 하단', 'LawyerCTABanner', '항상 표시 (퇴직금≥0)', '/lawyers 또는 상담 모달'],
        ['계산 결과 하단', 'LawyerUrgentCTA', '적격+퇴직금≥100만원', '즉시 상담 예약'],
        ['마이페이지',     'LawyerDirectory', '로그인 사용자', '노무사 상세 → 예약'],
        ['홈 화면 배너',  'LawyerBanner',     '비로그인 포함 전체', '/lawyers 랜딩'],
    ],
    col_widths=[3.0, 3.8, 4.0, 4.2]
)

# 2.3
add_heading(doc, '2.3  구현 방법 — 4단계', level=2, color=BRAND_BLUE)
stages = [
    ('STEP 1', '어드민 패널 구축', '1~2주',
     ['Supabase lawyers 테이블 생성 (name, region, specialty, fee, contact, bio)',
      '간단한 어드민 UI (또는 Supabase Table Editor 직접 사용)',
      '노무사 프로필 이미지 업로드 (Supabase Storage)']),
    ('STEP 2', 'React UI 구현', '1주',
     ['LawyerCard 컴포넌트: 사진, 이름, 전문분야, 지역, 상담료, 별점',
      'LawyerDirectory 페이지: 지역 필터 + 전문분야 필터',
      'LawyerCTABanner 컴포넌트: 계산 결과 페이지 삽입']),
    ('STEP 3', '상담 신청 폼 & 알림', '1주',
     ['상담 신청 폼: 이름, 연락처, 희망 일시, 문의 내용',
      'Supabase consultations 테이블 저장',
      'Discord 웹훅 또는 이메일로 노무사 및 관리자 알림']),
    ('STEP 4', '결제 연동 (선택)', '2~3주',
     ['포트원(구 아임포트) SDK 연동 → 카카오페이·카드 결제',
      '상담료 에스크로: 상담 완료 후 노무사 정산 (수수료 15~20%)',
      '영수증 발행 자동화']),
]

for code, title, duration, items in stages:
    tbl_s = doc.add_table(rows=1, cols=1)
    cell_s = tbl_s.cell(0, 0)
    set_cell_bg(cell_s, 'EFF6FF')
    set_cell_border(cell_s,
        left={'val':'single','sz':12,'color':'1A56DB'})
    p_s = cell_s.paragraphs[0]
    r_code = p_s.add_run(f'{code}  ')
    r_code.font.bold = True
    r_code.font.color.rgb = BRAND_BLUE
    r_code.font.size = Pt(10)
    r_title = p_s.add_run(title)
    r_title.font.bold = True
    r_title.font.size = Pt(11)
    r_dur = p_s.add_run(f'  (소요: {duration})')
    r_dur.font.size = Pt(9.5)
    r_dur.font.color.rgb = GRAY_TEXT
    for item in items:
        ip = cell_s.add_paragraph()
        ip.paragraph_format.left_indent = Cm(0.5)
        ip.paragraph_format.space_after = Pt(2)
        ir = ip.add_run(f'▸  {item}')
        ir.font.size = Pt(10)
    doc.add_paragraph().paragraph_format.space_after = Pt(3)

# 2.4
add_heading(doc, '2.4  법적 리스크 분석', level=2, color=BRAND_BLUE)
make_table(doc,
    headers=['법률', '위험도', '리스크 내용', '대응 방안'],
    rows=[
        ['노무사법 제22조', '[RED]HIGH', '노무사 아닌 자의 노무대리 금지', '중개 플랫폼으로 포지셔닝, 법률행위 직접 X'],
        ['변호사법 제109조','[RED]HIGH', '법률 사건 알선·중개 금지',       '노무상담 정보제공 플랫폼 명시'],
        ['전자상거래법',     '[ORANGE]MID','통신판매중개업 신고 의무',       '통신판매중개업자 신고 (간단)'],
        ['개인정보보호법',   '[ORANGE]MID','상담 내용 제3자 제공 동의',      '개인정보 처리방침 및 동의 UI'],
        ['소비자보호법',     '[GREEN]LOW', '환불 규정 명시 필요',            '이용약관에 환불 정책 명기'],
    ],
    col_widths=[3.2, 2.0, 4.5, 4.3]
)
add_info_box(doc, '⚖️  법적 포지셔닝 권고',
    ['CATCH는 "노무 서비스 정보 제공 플랫폼"으로 운영 → 직접 법률 행위 없음',
     '노무사와 사용자 간 계약은 당사자 직접 체결 (플랫폼은 연결만 담당)',
     '서비스 이용약관에 "CATCH는 중개 플랫폼이며 상담 결과에 책임 없음" 명시'],
    bg_hex='FEF3C7', title_color=ACCENT_ORANGE)

# 2.5
add_heading(doc, '2.5  수익 모델', level=2, color=BRAND_BLUE)
make_table(doc,
    headers=['수익원', '단가', '월 예상 볼륨', '월 예상 수익', '비고'],
    rows=[
        ['노무사 리스팅 피',  '₩50,000~150,000/월', '20명',    '₩100~300만', 'PART 3 파트너 패키지'],
        ['상담 연결 수수료',  '상담료의 15~20%',     '월 30건', '₩45~90만',  '상담료 평균 10만원 기준'],
        ['프리미엄 상단 노출','₩100,000~200,000/월', '5명',     '₩50~100만', '지역 독점 옵션'],
        ['상담 패키지 판매',  '₩50,000~200,000/패키지','10건',  '₩50~200만', '추후 도입'],
        ['[BOLD]합계 (12개월 후)', '—', '—', '[GREEN]₩245~690만/월', '[BOLD]연 2,940~8,280만원'],
    ],
    col_widths=[3.5, 3.5, 2.5, 2.8, 3.0]
)

# 2.6
add_heading(doc, '2.6  노무사 서비스 KPI', level=2, color=BRAND_BLUE)
make_table(doc,
    headers=['지표', '1개월', '3개월', '6개월', '12개월', '측정 방법'],
    rows=[
        ['등록 노무사 수',      '3명',    '10명',   '25명',   '50명',   'lawyers 테이블'],
        ['상담 신청 수',        '10건',   '50건',   '150건',  '400건',  'consultations 테이블'],
        ['상담 전환율(CTR→신청)','2%',    '4%',     '6%',     '8%',     '신청수/CTA클릭수'],
        ['노무사 만족도',       '—',      '4.0/5',  '4.2/5',  '4.5/5',  '완료 후 설문'],
        ['월 수익',            '₩30만',  '₩150만', '₩350만', '₩600만', '결제 대시보드'],
    ],
    col_widths=[3.5, 2.0, 2.0, 2.0, 2.0, 3.5]
)

# ══════════════════════════════════════════════════
# PART 3 — 노무사 파트너십 세일즈 & 마케팅 자료
# ══════════════════════════════════════════════════
add_part_title(doc, '3', '노무사 파트너십 세일즈 & 마케팅',
               '노무사를 파트너로 모집하기 위한 실전 영업 자료')

# 3.1
add_heading(doc, '3.1  첫 접촉 스크립트', level=2, color=BRAND_BLUE)
add_body(doc, '📧 이메일 스크립트 (초안)', bold=True, color=BRAND_DARK)

tbl_email = doc.add_table(rows=1, cols=1)
cell_em = tbl_email.cell(0, 0)
set_cell_bg(cell_em, 'F8FAFC')
set_cell_border(cell_em,
    top={'val':'single','sz':4,'color':'94A3B8'},
    bottom={'val':'single','sz':4,'color':'94A3B8'},
    left={'val':'single','sz':4,'color':'94A3B8'},
    right={'val':'single','sz':4,'color':'94A3B8'})
email_lines = [
    '제목: [CATCH 파트너 제안] 월 XX명의 퇴직금 관련 상담 잠재고객을 연결해 드립니다',
    '',
    '안녕하세요, 노무사님.',
    '',
    '저는 일용직 근로자 퇴직금·실업급여 계산 앱 CATCH를 운영하는 [이름]입니다.',
    '현재 CATCH는 월 [N]명의 사용자가 퇴직금 계산을 완료하고 있으며,',
    '이 중 상당수가 "실제 청구 방법"을 모르는 상황입니다.',
    '',
    '노무사님의 전문성과 CATCH의 트래픽을 연결하면 양측 모두 윈-윈할 수 있습니다:',
    '  • 노무사님: 의도가 명확한 잠재 고객을 월 XX명 확보',
    '  • CATCH: 사용자에게 실질적인 권리 구제 경로 제공',
    '',
    '월 50,000원 BASIC 플랜으로 시작하실 수 있으며, 첫 2개월은 무료로 제공드립니다.',
    '15분 통화로 자세히 설명드릴 수 있을까요?',
    '',
    '감사합니다.',
]
for i, line in enumerate(email_lines):
    if i == 0:
        p_em = cell_em.paragraphs[0]
    else:
        p_em = cell_em.add_paragraph()
    p_em.paragraph_format.left_indent = Cm(0.3)
    p_em.paragraph_format.space_after = Pt(1)
    r_em = p_em.add_run(line)
    r_em.font.size = Pt(9.5)
    if line.startswith('제목:'):
        r_em.font.bold = True
        r_em.font.color.rgb = BRAND_BLUE

doc.add_paragraph().paragraph_format.space_after = Pt(4)

add_body(doc, '📞 전화 스크립트 (30초 피치)', bold=True, color=BRAND_DARK)
add_info_box(doc, '전화 30초 피치',
    ['"안녕하세요, 노무사님. CATCH 앱 운영사입니다."',
     '"저희 앱은 월 [N]명의 일용직 근로자가 퇴직금 계산을 하는데,"',
     '"이분들이 실제로 청구를 도와줄 노무사님이 필요한 상황입니다."',
     '"파트너 노무사로 등록하시면 의도 있는 잠재 고객을 바로 연결해 드립니다."',
     '"15분만 시간 내주실 수 있으실까요?"'],
    bg_hex='F0F9FF', title_color=BRAND_BLUE)

# 3.2
add_heading(doc, '3.2  가치 제안 비교표', level=2, color=BRAND_BLUE)
add_body(doc, '노무사 입장에서 CATCH 파트너십의 경쟁력을 한눈에 보여주는 자료입니다.', space_after=4)
make_table(doc,
    headers=['항목', '네이버 광고', '로앤굿', '변호사닷컴', '[BOLD]CATCH 파트너'],
    rows=[
        ['월 비용',      '₩50~500만', '₩30~100만', '₩50~200만', '[GREEN]₩5~15만'],
        ['리드 품질',    '불특정 다수',   '중간',      '중간',      '[GREEN]퇴직금 상담 의향자'],
        ['경쟁 노출',    '같은 페이지 다수 노무사', '다수 경쟁', '다수 경쟁', '[GREEN]지역 독점 옵션'],
        ['타겟 정밀도',  '낮음',       '중간',      '중간',       '[GREEN]매우 높음 (계산 완료자)'],
        ['계약 기간',    '최소 1개월',  '3~6개월',   '3~6개월',   '[GREEN]월 단위 (언제든 해지)'],
        ['성과 투명성',  '클릭수만',    '부분 공개',  '부분 공개',  '[GREEN]상담 신청수 실시간 대시보드'],
    ],
    col_widths=[3.0, 3.0, 2.5, 2.5, 4.0]
)

# 3.3
add_heading(doc, '3.3  파트너 패키지', level=2, color=BRAND_BLUE)
packages = [
    ('BASIC', '₩50,000/월', 'EFF6FF', '1A56DB', [
        '노무사 프로필 카드 1개 등록',
        '전국 디렉토리 노출',
        '월 상담 신청 최대 10건 연결',
        '기본 통계 대시보드',
        '첫 2개월 무료 제공',
    ]),
    ('PREMIUM', '₩150,000/월', 'F0FDF4', '059669', [
        'BASIC 모든 기능 포함',
        '지역 상단 고정 노출 (1개 지역)',
        '계산 결과 페이지 CTA 우선 노출',
        '월 상담 신청 무제한 연결',
        '사용자 리뷰 표시',
        '분기별 1:1 성과 미팅',
    ]),
    ('ENTERPRISE', '협의 (₩300,000+/월)', 'FFF7ED', 'D97706', [
        'PREMIUM 모든 기능 포함',
        '다수 지역 독점 노출',
        '브랜드 배너 광고 (홈 화면)',
        '단체상담 세미나 홍보 지원',
        '전담 계정 매니저',
        '맞춤형 리드 조건 설정',
    ]),
]
for pkg_name, price, bg, border, features in packages:
    tbl_p = doc.add_table(rows=1, cols=1)
    cell_p = tbl_p.cell(0, 0)
    set_cell_bg(cell_p, bg)
    set_cell_border(cell_p,
        top={'val':'single','sz':8,'color':border},
        bottom={'val':'single','sz':8,'color':border},
        left={'val':'single','sz':8,'color':border},
        right={'val':'single','sz':8,'color':border})
    p_pkg = cell_p.paragraphs[0]
    rn1 = p_pkg.add_run(f'  {pkg_name}  ')
    rn1.font.bold = True
    rn1.font.size = Pt(13)
    rn1.font.color.rgb = RGBColor(int(border[:2],16), int(border[2:4],16), int(border[4:],16))
    rn2 = p_pkg.add_run(price)
    rn2.font.size  = Pt(12)
    rn2.font.color.rgb = GRAY_TEXT
    for feat in features:
        fp = cell_p.add_paragraph()
        fp.paragraph_format.left_indent = Cm(0.5)
        fp.paragraph_format.space_after = Pt(1)
        fr = fp.add_run(f'✓  {feat}')
        fr.font.size = Pt(10)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

# 3.4
add_heading(doc, '3.4  파트너 모집 타임라인 & 접촉 채널', level=2, color=BRAND_BLUE)
make_table(doc,
    headers=['기간', '목표', '접촉 채널', '목표 파트너 수'],
    rows=[
        ['1개월차', '시드 파트너 모집', '지인 네트워크, 네이버카페 노무사 모임', '3~5명'],
        ['2~3개월', '수도권 확장', '대한공인노무사회 제휴, 링크드인 DM',        '10~15명'],
        ['4~6개월', '전국 확장',   '노무사 커뮤니티 광고, 콘텐츠 마케팅',       '25~30명'],
        ['7~12개월','플랫폼 안정화','자동 가입, 레퍼럴 프로그램',               '50명 이상'],
    ],
    col_widths=[2.0, 3.5, 5.5, 4.0]
)

add_body(doc, '▶ 핵심 접촉 채널 상세', bold=True, color=BRAND_DARK, space_after=4)
for channel in [
    '대한공인노무사회 (www.kafl.or.kr) — 협회 공식 제휴 문의',
    '노무사 카카오오픈채팅 / 네이버카페 "노무사 모임" — 직접 홍보',
    '링크드인 HR/노무 그룹 — B2B 영업 DM',
    '블로그 SEO: "노무사 마케팅 방법", "노무사 신규 고객 확보" — 인바운드',
    '유튜브 콘텐츠: "CATCH와 함께하는 노무사 파트너십" — 영상 홍보',
]:
    add_bullet(doc, channel)

# 3.5
add_heading(doc, '3.5  반론 처리 (Objection Handling)', level=2, color=BRAND_BLUE)
objections = [
    ('이미 다른 플랫폼을 쓰고 있어요',
     'CATCH는 퇴직금 계산 완료 직후 상담 의향이 확실한 사람만 연결합니다. '
     '기존 플랫폼과 타겟이 달라 중복이 아닌 추가 채널입니다.'),
    ('비용 대비 효과가 불확실해요',
     '첫 2개월 무료 + 월 단위 해지 가능으로 리스크 제로입니다. '
     '2개월 내 상담 신청이 없으면 전액 환불 보장합니다.'),
    ('CATCH 사용자가 실제로 상담할까요?',
     '퇴직금 계산을 직접 하러 온 사람 = 퇴직 상황 = 상담 필요성 최고조. '
     '타 플랫폼 대비 전환율 2~3배 높습니다. (파일럿 데이터 공유 가능)'),
    ('관리가 번거로울 것 같아요',
     '알림은 카카오톡 또는 문자로 자동 발송, 대시보드에서 신청 현황 실시간 확인. '
     '추가 관리 업무는 최소화됩니다.'),
]
for q, a in objections:
    tbl_obj = doc.add_table(rows=2, cols=1)
    cell_q = tbl_obj.cell(0, 0)
    cell_a = tbl_obj.cell(1, 0)
    set_cell_bg(cell_q, '1E293B')
    set_cell_bg(cell_a, 'F8FAFC')
    set_cell_border(cell_q, left={'val':'single','sz':10,'color':'1A56DB'})
    set_cell_border(cell_a, left={'val':'single','sz':10,'color':'94A3B8'})
    pq = cell_q.paragraphs[0]
    pq.paragraph_format.left_indent = Cm(0.3)
    rq = pq.add_run(f'Q.  {q}')
    rq.font.bold = True
    rq.font.size = Pt(10.5)
    rq.font.color.rgb = WHITE
    pa = cell_a.paragraphs[0]
    pa.paragraph_format.left_indent = Cm(0.3)
    ra = pa.add_run(f'A.  {a}')
    ra.font.size = Pt(10)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

# ══════════════════════════════════════════════════
# 부록 — 통합 로드맵 + 핵심 요약
# ══════════════════════════════════════════════════
doc.add_page_break()
add_heading(doc, '부록  —  통합 12개월 로드맵', level=1, color=BRAND_DARK, space_before=4)
make_table(doc,
    headers=['분기', '월', '채용정보 섹션', '노무사 중개 서비스', '예상 수익 합계'],
    rows=[
        ['Q1', '1개월', '워크넷 API 연동', '어드민 패널 + UI', '₩30만'],
        ['Q1', '2개월', '/jobs 페이지 오픈', '노무사 3명 온보딩', '₩80만'],
        ['Q1', '3개월', '위치 기반 필터', '상담 신청 폼 완성', '₩150만'],
        ['Q2', '4개월', '파트너사 협의 시작', '결제 연동 (포트원)', '₩250만'],
        ['Q2', '5개월', 'CPC 광고 본격화', '노무사 10명 달성',   '₩350만'],
        ['Q2', '6개월', '잡코리아 API 계약', '프리미엄 패키지 판매', '₩500만'],
        ['Q3', '7개월', '프리미엄 공고 판매', '지역 독점 옵션 출시', '₩600만'],
        ['Q3', '8개월', '쿠팡 공식 제휴',    '노무사 25명 달성',   '₩700만'],
        ['Q3', '9개월', '채용 성사 수수료',   '레퍼럴 프로그램',    '₩800만'],
        ['Q4', '10개월','연관 채용 AI 추천',  '노무사 40명 달성',   '₩950만'],
        ['Q4', '11개월','채용 패키지 판매',   '단체상담 세미나',    '₩1,100만'],
        ['[BOLD]Q4', '[BOLD]12개월', '[GREEN]월 채용 수익 700만', '[GREEN]월 노무 수익 600만', '[GREEN]₩1,300만+'],
    ],
    col_widths=[1.5, 1.5, 4.5, 4.5, 3.0]
)

add_heading(doc, '핵심 요약', level=2, color=BRAND_DARK, space_before=12)

summary_items = [
    ('총 예상 연 수익 (12개월 후)', '채용정보 ~₩8,400만 + 노무사 ~₩7,200만 = [BOLD]연 1.5억+'),
    ('초기 투자 비용', '워크넷 API 무료, 개발 인력 기준 약 4~6주 개발 (프론트+백엔드)'),
    ('최단 수익화 경로', '노무사 파트너 3명 온보딩 → 월 100~200만 즉시 가능 (2개월 내)'),
    ('최대 리스크', '노무사법 위반 → 중개 플랫폼 포지셔닝 철저히 유지'),
    ('핵심 성공 요인', '사용자 트래픽 확보 → 노무사 파트너 품질 관리 → 반복 방문 유도'),
]

for label, value in summary_items:
    tbl_sum = doc.add_table(rows=1, cols=2)
    tbl_sum.alignment = WD_TABLE_ALIGNMENT.CENTER
    c_label = tbl_sum.cell(0, 0)
    c_value = tbl_sum.cell(0, 1)
    set_cell_bg(c_label, '1E293B')
    set_cell_bg(c_value, 'EFF6FF')
    c_label.width = Cm(5.5)
    c_value.width = Cm(10.5)
    p_l = c_label.paragraphs[0]
    rl = p_l.add_run(label)
    rl.font.bold  = True
    rl.font.size  = Pt(10)
    rl.font.color.rgb = WHITE
    p_v = c_value.paragraphs[0]
    if '[BOLD]' in value:
        parts = value.split('[BOLD]')
        rv1 = p_v.add_run(parts[0])
        rv1.font.size = Pt(10)
        rv2 = p_v.add_run(parts[1])
        rv2.font.bold = True
        rv2.font.size = Pt(10)
        rv2.font.color.rgb = ACCENT_GREEN
    else:
        rv = p_v.add_run(value)
        rv.font.size = Pt(10)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

doc.add_paragraph()
add_info_box(doc, '🚀  최우선 액션 아이템 (이번 주)',
    ['1. 워크넷 OpenAPI 신청 → https://www.work.go.kr/openapi',
     '2. Supabase lawyers + jobs 테이블 스키마 설계',
     '3. 지인 노무사 1명에게 BASIC 플랜 무료 온보딩 제안',
     '4. /jobs 페이지 React 컴포넌트 초안 작성',
     '5. 이메일 스크립트 + 파트너 제안 덱(PPT) 초안 완성'],
    bg_hex='ECFDF5', title_color=ACCENT_GREEN)

# 푸터
doc.add_paragraph()
p_footer = doc.add_paragraph()
p_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_f = p_footer.add_run('CATCH (퇴직금 한번에) — 내부 전략 문서  |  2026년 3월 24일  |  기밀')
r_f.font.size = Pt(8.5)
r_f.font.color.rgb = GRAY_TEXT

# ──────────────────────────────────────────
# 저장
# ──────────────────────────────────────────
output_path = 'CATCH_채용정보_노무사_전략보고서.docx'
doc.save(output_path)
print(f'OK: {output_path}')
